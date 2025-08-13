import azure.functions as func
import os
import json
import uuid
import traceback
from datetime import datetime, timedelta
from docx import Document
from docx.shared import Pt
import requests
import re
from io import BytesIO
from azure.storage.blob import BlobServiceClient, generate_blob_sas, BlobSasPermissions

# Configuración
AZURE_OPENAI_ENDPOINT = os.getenv("AZURE_OPENAI_ENDPOINT", "https://chabot-inventario-talento-aistudio.openai.azure.com/")
AZURE_OPENAI_API_KEY = os.getenv("AZURE_OPENAI_API_KEY")
DEPLOYMENT_NAME = os.getenv("DEPLOYMENT_NAME", "gpt-4o-mini")
API_VERSION = "2024-02-15-preview"

STORAGE_CONNECTION_STRING = os.getenv("STORAGE_CONNECTION_STRING")
BLOB_CONTAINER_NAME = "propia"
TEMPLATE_CONTAINER_NAME = "propia"
TEMPLATE_BLOB_NAME = "plantilla/Plantilla-Propuesta.docx"
PROPUESTAS_FOLDER = "propuestas/"

# Registrar la función
app = func.FunctionApp()

# ========== FUNCIONES AUXILIARES ==========

def call_azure_openai(messages, max_tokens=1000):
    """Función para llamar a Azure OpenAI"""
    try:
        if not AZURE_OPENAI_API_KEY:
            raise Exception("Azure OpenAI API Key no configurado")
            
        api_url = f"{AZURE_OPENAI_ENDPOINT}openai/deployments/{DEPLOYMENT_NAME}/chat/completions?api-version={API_VERSION}"

        headers = {
            "Content-Type": "application/json",
            "api-key": AZURE_OPENAI_API_KEY
        }

        data = {
            "messages": messages,
            "max_tokens": max_tokens
        }

        response = requests.post(api_url, headers=headers, json=data)

        if response.status_code == 200:
            return response.json()['choices'][0]['message']['content'].strip()
        else:
            raise Exception(f"Error Azure OpenAI: {response.status_code}, {response.text}")

    except Exception as e:
        raise Exception(f"Error llamando Azure OpenAI: {str(e)}")

def limpiar_formato_markdown(texto):
    """Limpia formato Markdown para Word"""
    if not texto:
        return texto
        
    # Convertir tablas markdown a formato de texto estructurado
    tabla_pattern = r'\|([^|]+\|)+\n\|[-:|\s]+\|\n(\|([^|]+\|)+\n?)+'
    tablas = re.findall(tabla_pattern, texto)
    
    for i, tabla in enumerate(tablas):
        lineas_tabla = tabla[0].split('\n')
        texto_tabla_formateado = ""
        
        for linea in lineas_tabla:
            if '|' in linea and not re.match(r'\|[-:\s|]+\|', linea):
                celdas = [celda.strip() for celda in linea.split('|') if celda.strip()]
                if celdas:
                    texto_tabla_formateado += " • " + " - ".join(celdas) + "\n"
        
        texto = texto.replace(tabla[0], texto_tabla_formateado)
    
    # Limpiar formato markdown
    texto = re.sub(r'^#{1,6}\s+', '', texto, flags=re.MULTILINE)
    texto = re.sub(r'\*\*(.*?)\*\*', r'\1', texto)
    texto = re.sub(r'\*(.*?)\*', r'\1', texto)
    texto = re.sub(r'^\s*[-\*\+]\s+', '• ', texto, flags=re.MULTILINE)
    texto = re.sub(r'```(?:\w+)?\n?(.*?)\n?```', r'\1', texto, flags=re.DOTALL)
    texto = re.sub(r'\n\s*\n\s*\n', '\n\n', texto)
    
    return texto.strip()

def extraer_informacion_empresa(prompt_completo):
    """Extrae información específica de la empresa del prompt"""
    empresa_pattern = r'(?:para\s+|de\s+|empresa\s+)([A-Z][a-zA-Z\s&]+(?:SA\s+de\s+CV|S\.A\.|Inc\.|Corp\.|Ltd\.)?)'
    fecha_pattern = r'(\d{1,2}\s+de\s+\w+\s+de\s+\d{4}|\d{1,2}/\d{1,2}/\d{4})'
    titulo_pattern = r'#\s*(.*?)(?:\n|$)'
    
    info = {
        'empresa': 'Cliente Estimado',
        'fecha': datetime.now().strftime('%d de %B de %Y'),
        'titulo': 'Propuesta Técnica'
    }
    
    empresa_match = re.search(empresa_pattern, prompt_completo, re.IGNORECASE)
    if empresa_match:
        info['empresa'] = empresa_match.group(1).strip()
    
    fecha_match = re.search(fecha_pattern, prompt_completo)
    if fecha_match:
        info['fecha'] = fecha_match.group(1)
    
    titulo_match = re.search(titulo_pattern, prompt_completo)
    if titulo_match:
        info['titulo'] = titulo_match.group(1).strip()
    
    return info

# ========== FUNCIONES DE GENERACIÓN DE CONTENIDO ==========

def generar_resumen_ejecutivo(prompt_completo):
    """Genera contenido específico para resumen ejecutivo"""
    messages = [
        {
            "role": "system",
            "content": """Eres un consultor experto en propuestas técnicas. Genera ÚNICAMENTE un resumen ejecutivo profesional para documentos Word.

IMPORTANTE:
- NO uses formato Markdown (sin #, **, -, etc.)
- Escribe en párrafos corridos para Word
- Estilo profesional y ejecutivo
- Máximo 4 párrafos
- No incluyas títulos ni encabezados
- Enfócate solo en el resumen ejecutivo del proyecto"""
        },
        {
            "role": "user",
            "content": f"Basándote en esta información, genera únicamente el resumen ejecutivo profesional:\n\n{prompt_completo}"
        }
    ]
    
    contenido = call_azure_openai(messages, max_tokens=600)
    return limpiar_formato_markdown(contenido) if contenido else None

def generar_alcance_minimo(prompt_completo):
    """Genera contenido específico para alcance mínimo"""
    messages = [
        {
            "role": "system",
            "content": """Eres un consultor experto en propuestas técnicas. Genera ÚNICAMENTE el alcance mínimo del proyecto para documentos Word.

IMPORTANTE:
- NO uses formato Markdown (sin #, **, -, etc.)
- Escribe en párrafos corridos para Word
- Describe qué incluye el proyecto específicamente
- Máximo 5 párrafos
- No incluyas títulos ni encabezados
- Enfócate solo en el alcance del proyecto"""
        },
        {
            "role": "user",
            "content": f"Basándote en esta información, genera únicamente el alcance mínimo del proyecto:\n\n{prompt_completo}"
        }
    ]
    
    contenido = call_azure_openai(messages, max_tokens=700)
    return limpiar_formato_markdown(contenido) if contenido else None

def generar_plan_trabajo(prompt_completo):
    """Genera contenido específico para plan de trabajo"""
    messages = [
        {
            "role": "system",
            "content": """Eres un consultor experto en propuestas técnicas. Genera ÚNICAMENTE el plan de trabajo del proyecto para documentos Word.

IMPORTANTE:
- Si hay información tabular, preséntala usando bullets (•) en formato estructurado
- Conserva TODOS los números, fechas, porcentajes y datos cuantitativos
- Si hay fases con duraciones, inclúyelas con sus tiempos específicos
- Usa formato: "• Fase X: Descripción - Duración: X semanas - Porcentaje: X%"
- NO uses formato Markdown tabla (|---|) 
- Escribe en párrafos y listas con bullets para Word
- Máximo 6 párrafos o secciones con bullets
- Incluye TODA la información numérica disponible"""
        },
        {
            "role": "user",
            "content": f"Basándote en esta información, genera la descripción del plan de trabajo incluyendo TODAS las fases, duraciones y porcentajes mostrados:\n\n{prompt_completo}"
        }
    ]
    
    contenido = call_azure_openai(messages, max_tokens=800)
    return limpiar_formato_markdown(contenido) if contenido else None

def generar_estructura_equipo(prompt_completo):
    """Genera contenido específico para estructura del equipo"""
    messages = [
        {
            "role": "system",
            "content": """Eres un consultor experto en propuestas técnicas. Genera ÚNICAMENTE la descripción del equipo de trabajo para documentos Word.

IMPORTANTE:
- Si hay información de roles con costos, preséntala usando bullets (•)
- Conserva TODOS los números, tarifas, horas y subtotales
- Usa formato: "• Rol: Descripción - Dedicación: X% - Horas: X - Tarifa: $X - Subtotal: $X"
- NO uses formato Markdown tabla (|---|)
- Incluye descuentos y totales si están disponibles
- Escribe en párrafos y listas con bullets para Word
- Máximo 5 párrafos o secciones con bullets
- Incluye TODA la información numérica y financiera disponible"""
        },
        {
            "role": "user",
            "content": f"Basándote en esta información, genera la descripción del equipo incluyendo TODOS los roles, costos, tarifas y totales mostrados:\n\n{prompt_completo}"
        }
    ]
    
    contenido = call_azure_openai(messages, max_tokens=800)
    return limpiar_formato_markdown(contenido) if contenido else None

def generar_inversion_detallada(prompt_completo):
    """Genera contenido específico para inversión detallada"""
    messages = [
        {
            "role": "system",
            "content": """Eres un consultor experto en propuestas técnicas. Genera ÚNICAMENTE la explicación de la inversión detallada para documentos Word.

IMPORTANTE:
- Conserva TODOS los números, montos, porcentajes y cifras exactas
- Si hay desglose de costos, preséntalo usando bullets (•)
- Usa formato: "• Concepto: Descripción - Monto: $X,XXX MXN"
- NO uses formato Markdown tabla (|---|)
- Incluye servicios profesionales, costos de setup, costos operativos
- Menciona la inversión total inicial con el monto exacto
- Escribe en párrafos y listas con bullets para Word
- Máximo 4 párrafos o secciones con bullets
- Incluye TODA la información financiera disponible"""
        },
        {
            "role": "user",
            "content": f"Basándote en esta información, genera la explicación de la inversión incluyendo TODOS los montos, costos y totales mostrados:\n\n{prompt_completo}"
        }
    ]
    
    contenido = call_azure_openai(messages, max_tokens=700)
    return limpiar_formato_markdown(contenido) if contenido else None

def generar_supuestos_condiciones(prompt_completo):
    """Genera contenido específico para supuestos y condiciones"""
    messages = [
        {
            "role": "system",
            "content": """Eres un consultor experto en propuestas técnicas. Genera ÚNICAMENTE los supuestos y condiciones del proyecto para documentos Word.

IMPORTANTE:
- NO uses formato Markdown (sin #, **, -, etc.)
- Escribe en párrafos corridos para Word
- Describe supuestos técnicos y condiciones comerciales
- Máximo 4 párrafos
- No incluyas títulos ni encabezados
- Enfócate en aspectos clave del proyecto"""
        },
        {
            "role": "user",
            "content": f"Basándote en esta información, genera únicamente los supuestos y condiciones:\n\n{prompt_completo}"
        }
    ]
    
    contenido = call_azure_openai(messages, max_tokens=700)
    return limpiar_formato_markdown(contenido) if contenido else None

def generar_carta_presentacion(prompt_completo):
    """Genera contenido específico para carta de presentación"""
    info_empresa = extraer_informacion_empresa(prompt_completo)
    
    messages = [
        {
            "role": "system",
            "content": f"""Eres un consultor experto en propuestas técnicas. Genera ÚNICAMENTE una carta de presentación profesional para documentos Word.

IMPORTANTE:
- NO uses formato Markdown (sin #, **, -, etc.)
- Escribe en párrafos corridos para Word
- Usa un tono profesional y cordial
- Menciona específicamente el proyecto y sus características principales
- Incluye los datos extraídos: Empresa: {info_empresa['empresa']}, Fecha: {info_empresa['fecha']}
- Máximo 3 párrafos
- No incluyas títulos ni encabezados
- Enfócate en el valor que HITSS puede aportar al proyecto"""
        },
        {
            "role": "user",
            "content": f"Genera una carta de presentación profesional para la empresa {info_empresa['empresa']} basándote en este proyecto:\n\n{prompt_completo}"
        }
    ]
    
    contenido = call_azure_openai(messages, max_tokens=500)
    
    if contenido:
        carta_personalizada = f"""Estimado Equipo,

Ciudad de México, México {info_empresa['fecha']}

En representación de HITSS, agradecemos profundamente la oportunidad que nos brindan de presentar nuestra propuesta para {info_empresa['titulo'].lower()}. {contenido}

Atentamente,
Sergio Portales Aburto"""
        
        return limpiar_formato_markdown(carta_personalizada)
    
    return None

def generar_titulo_fecha(prompt_completo):
    """Genera título y fecha para el cuadro de texto de la página 1"""
    info_empresa = extraer_informacion_empresa(prompt_completo)
    
    titulo_fecha_content = f"""{info_empresa['titulo']}
{info_empresa['fecha']}"""
    
    return titulo_fecha_content

# ========== FUNCIONES DE PROCESAMIENTO DE DOCUMENTOS ==========

def replace_in_paragraph(paragraph, old_text, new_text, font_name="Arial Nova Cond", font_size=11.5):
    """Reemplaza texto en párrafos"""
    if old_text in paragraph.text:
        full_text = paragraph.text.replace(old_text, new_text)
        
        for run in paragraph.runs:
            run.clear()
        
        new_run = paragraph.add_run(full_text)
        new_run.font.name = font_name
        new_run.font.size = Pt(font_size)
        
        return True
    return False

def replace_in_tables(doc, old_text, new_text, font_name="Arial Nova Cond", font_size=12):
    """Reemplaza texto en tablas del documento"""
    replacements_made = 0
    
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if replace_in_paragraph(paragraph, old_text, new_text, font_name, font_size):
                        replacements_made += 1
    
    return replacements_made

def replace_in_textboxes(doc, old_text, new_text):
    """Reemplaza texto en cuadros de texto usando el método que funciona"""
    replacements_made = 0
    
    try:
        # Buscar elementos p que contienen exactamente el placeholder
        for element in doc.element.iter():
            if element.tag.endswith('}p'):
                # Obtener el texto completo del párrafo concatenando todos los elementos t
                full_text = ""
                text_elements = []
                
                for child in element.iter():
                    if child.tag.endswith('}t') and child.text:
                        full_text += child.text
                        text_elements.append(child)
                
                # Si el texto completo contiene el placeholder
                if old_text in full_text:
                    # Verificar si está en txbxContent (textbox content)
                    parent = element.getparent()
                    if parent is not None and 'txbxContent' in parent.tag:
                        # Limpiar todos los elementos de texto existentes
                        for text_elem in text_elements:
                            text_elem.text = ""
                        
                        # Poner el texto reemplazado en el primer elemento
                        if text_elements:
                            new_text_content = full_text.replace(old_text, new_text)
                            text_elements[0].text = new_text_content
                            replacements_made += 1
        
        # Método de respaldo: usar el método original si no se encontraron reemplazos
        if replacements_made == 0:
            for element in doc.element.body.iter():
                for run_element in element.iter():
                    if run_element.tag.endswith('}r'):
                        for text_element in run_element.iter():
                            if text_element.tag.endswith('}t') and text_element.text:
                                if old_text in text_element.text:
                                    text_element.text = text_element.text.replace(old_text, new_text)
                                    replacements_made += 1
        
    except Exception as e:
        raise Exception(f"Error procesando cuadros de texto: {str(e)}")
    
    return replacements_made

# ========== FUNCIONES DE AZURE STORAGE ==========

def get_blob_service_client():
    """Obtiene el cliente de Azure Blob Storage"""
    if not STORAGE_CONNECTION_STRING:
        raise Exception("Storage connection string no configurado")
    return BlobServiceClient.from_connection_string(STORAGE_CONNECTION_STRING)

def descargar_plantilla():
    """Descarga la plantilla desde Blob Storage"""
    try:
        blob_service_client = get_blob_service_client()
        blob_client = blob_service_client.get_blob_client(
            container=TEMPLATE_CONTAINER_NAME, 
            blob=TEMPLATE_BLOB_NAME
        )
        
        # Descargar el blob en memoria
        plantilla_stream = BytesIO()
        blob_data = blob_client.download_blob()
        plantilla_stream.write(blob_data.readall())
        plantilla_stream.seek(0)  # Regresar al inicio del stream para ser leído
        
        return plantilla_stream
    except Exception as e:
        raise Exception(f"Error descargando plantilla: {str(e)}")

def subir_documento(documento_stream, nombre_archivo):
    """Sube el documento generado a Blob Storage"""
    try:
        blob_service_client = get_blob_service_client()
        # Agregar el prefijo de carpeta propuestas/
        blob_name = f"{PROPUESTAS_FOLDER}{nombre_archivo}"
        
        blob_client = blob_service_client.get_blob_client(
            container=BLOB_CONTAINER_NAME,
            blob=blob_name
        )
        
        documento_stream.seek(0)
        blob_client.upload_blob(documento_stream, overwrite=True)
        
        return blob_name  # Retornar solo el nombre del blob para generar SAS después
    except Exception as e:
        raise Exception(f"Error subiendo documento: {str(e)}")

def generar_url_presignada(nombre_archivo, expiracion_minutos=60):
    """
    Genera una URL pre-firmada (SAS) para acceder al archivo en Azure Blob Storage.

    :param nombre_archivo: Nombre del archivo en el contenedor (puede incluir carpeta/).
    :param expiracion_minutos: Tiempo de expiración de la URL en minutos.
    :return: URL pre-firmada (SAS).
    """
    try:
        blob_service_client = get_blob_service_client()
        blob_client = blob_service_client.get_blob_client(
            container=BLOB_CONTAINER_NAME, 
            blob=nombre_archivo
        )

        # Configurar permisos y tiempo de expiración
        sas_token = generate_blob_sas(
            account_name=blob_client.account_name,
            container_name=BLOB_CONTAINER_NAME,
            blob_name=nombre_archivo,
            account_key=blob_service_client.credential.account_key,
            permission=BlobSasPermissions(read=True),  # Solo lectura
            expiry=datetime.utcnow() + timedelta(minutes=expiracion_minutos)  # Expira en X minutos
        )

        # Crear URL con el token SAS
        url_presignada = f"https://{blob_client.account_name}.blob.core.windows.net/{BLOB_CONTAINER_NAME}/{nombre_archivo}?{sas_token}"
        return url_presignada
    except Exception as e:
        raise Exception(f"Error generando URL pre-firmada: {str(e)}")

def verificar_documento_existe(nombre_archivo):
    """Verifica si un documento existe en Blob Storage"""
    try:
        blob_service_client = get_blob_service_client()
        # Agregar el prefijo de carpeta propuestas/ si no lo tiene
        if not nombre_archivo.startswith(PROPUESTAS_FOLDER):
            blob_name = f"{PROPUESTAS_FOLDER}{nombre_archivo}"
        else:
            blob_name = nombre_archivo
        
        blob_client = blob_service_client.get_blob_client(
            container=BLOB_CONTAINER_NAME,
            blob=blob_name
        )
        
        return blob_client.exists()
    except Exception as e:
        return False

def obtener_url_documento(nombre_archivo):
    """Obtiene la URL pre-firmada del documento en Blob Storage"""
    try:
        if not nombre_archivo.startswith(PROPUESTAS_FOLDER):
            blob_name = f"{PROPUESTAS_FOLDER}{nombre_archivo}"
        else:
            blob_name = nombre_archivo
            
        if verificar_documento_existe(nombre_archivo):
            return generar_url_presignada(blob_name, expiracion_minutos=120)  # 2 horas
        else:
            return None
    except Exception as e:
        raise Exception(f"Error obteniendo URL: {str(e)}")

# ========== FUNCIÓN PRINCIPAL DE PROCESAMIENTO ==========

def procesar_propuesta_completa(prompt_completo, document_id):
    """Procesa una propuesta completa"""
    try:
        # Descargar plantilla
        plantilla_stream = descargar_plantilla()
        doc = Document(plantilla_stream)
        
        # Extraer información de la empresa
        info_empresa = extraer_informacion_empresa(prompt_completo)
        
        # Definir placeholders con funciones de generación
        placeholders_config = {
            "[RESUMEN]": generar_resumen_ejecutivo,
            "[ALCANCE]": generar_alcance_minimo,
            "[PLAN_TRABAJO]": generar_plan_trabajo,
            "[EQUIPO]": generar_estructura_equipo,
            "[INVERSION]": generar_inversion_detallada,
            "[SUPUESTOS]": generar_supuestos_condiciones,
            "[CARTA_PRESENTACION]": generar_carta_presentacion,
            "[titulo]": lambda prompt: generar_titulo_fecha(prompt).split('\n')[0],
            "[fecha]": lambda prompt: generar_titulo_fecha(prompt).split('\n')[1]
        }
        
        cambios_totales = 0
        
        # Procesar cada placeholder
        for placeholder, funcion_generadora in placeholders_config.items():
            try:
                # Generar contenido específico
                contenido_generado = funcion_generadora(prompt_completo)
                
                if not contenido_generado:
                    continue
                
                replacements_for_this_item = 0
                
                # 1. Buscar en párrafos normales
                for paragraph in doc.paragraphs:
                    if replace_in_paragraph(paragraph, placeholder, contenido_generado):
                        replacements_for_this_item += 1
                
                # 2. Buscar en tablas
                table_replacements = replace_in_tables(doc, placeholder, contenido_generado)
                replacements_for_this_item += table_replacements
                
                # 3. Buscar en cuadros de texto
                textbox_replacements = replace_in_textboxes(doc, placeholder, contenido_generado)
                replacements_for_this_item += textbox_replacements
                
                cambios_totales += replacements_for_this_item
                
            except Exception as e:
                raise Exception(f"Error procesando {placeholder}: {str(e)}")
        
        if cambios_totales == 0:
            raise Exception("No se realizaron cambios en el documento")
        
        # Generar nombre de archivo
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        empresa_clean = re.sub(r'[^\w\s-]', '', info_empresa['empresa']).strip()[:20]
        nombre_archivo = f"Propuesta_{empresa_clean}_{document_id}_{timestamp}.docx"
        
        # Guardar documento en memoria
        documento_stream = BytesIO()
        doc.save(documento_stream)
        
        # Subir a Azure Storage
        blob_name = subir_documento(documento_stream, nombre_archivo)
        
        # Generar URL pre-firmada con expiración de 24 horas para el documento recién creado
        url_presignada = generar_url_presignada(blob_name, expiracion_minutos=1440)  # 24 horas
        
        return {
            "document_id": document_id,
            "filename": nombre_archivo,
            "blob_name": blob_name,
            "url_presignada": url_presignada,
            "empresa": info_empresa['empresa'],
            "fecha": info_empresa['fecha'],
            "titulo": info_empresa['titulo'],
            "cambios_realizados": cambios_totales,
            "status": "completed"
        }
        
    except Exception as e:
        raise Exception(f"Error procesando propuesta: {str(e)}")

# ========== AZURE FUNCTIONS ==========

@app.function_name(name="generar_propuesta")
@app.route(route="generar_propuesta", methods=["POST"], auth_level=func.AuthLevel.ANONYMOUS)
def generar_propuesta(req: func.HttpRequest) -> func.HttpResponse:
    """Endpoint POST para generar propuesta"""
    try:
        # Obtener el prompt del cuerpo de la solicitud
        try:
            req_body = req.get_json()
            if req_body and 'prompt' in req_body:
                prompt_completo = req_body['prompt']
            else:
                prompt_completo = req.get_body().decode('utf-8')
        except Exception:
            prompt_completo = req.get_body().decode('utf-8')
        
        if not prompt_completo or len(prompt_completo.strip()) == 0:
            return func.HttpResponse(
                json.dumps({
                    "error": "Prompt no proporcionado o vacío",
                    "document_id": None
                }),
                status_code=400,
                mimetype="application/json"
            )
        
        # Generar ID único para el documento
        document_id = str(uuid.uuid4())[:8]
        
        try:
            # Procesar la propuesta
            resultado = procesar_propuesta_completa(prompt_completo, document_id)
            
            return func.HttpResponse(
                json.dumps({
                    "message": "Propuesta generada exitosamente",
                    "document_id": document_id,
                    "filename": resultado["filename"],
                    "url_presignada": resultado["url_presignada"],
                    "expira_en_horas": 24,
                    "empresa": resultado["empresa"],
                    "titulo": resultado["titulo"],
                    "cambios_realizados": resultado["cambios_realizados"],
                    "status": "completed"
                }),
                status_code=200,
                mimetype="application/json"
            )
            
        except Exception as processing_error:
            return func.HttpResponse(
                json.dumps({
                    "error": f"Error procesando propuesta: {str(processing_error)}",
                    "document_id": document_id,
                    "status": "failed"
                }),
                status_code=500,
                mimetype="application/json"
            )
    
    except Exception as e:
        return func.HttpResponse(
            json.dumps({
                "error": f"Error interno del servidor: {str(e)}",
                "traceback": traceback.format_exc()
            }),
            status_code=500,
            mimetype="application/json"
        )

@app.function_name(name="obtener_propuesta")
@app.route(route="obtener_propuesta/{document_id}", methods=["GET"], auth_level=func.AuthLevel.ANONYMOUS)
def obtener_propuesta(req: func.HttpRequest) -> func.HttpResponse:
    """Endpoint GET para obtener URL del documento generado"""
    try:
        document_id = req.route_params.get('document_id')
        
        if not document_id:
            return func.HttpResponse(
                json.dumps({
                    "error": "Document ID no proporcionado",
                    "document_id": None
                }),
                status_code=400,
                mimetype="application/json"
            )
        
        try:
            # Buscar documentos que contengan el document_id en la carpeta propuestas/
            blob_service_client = get_blob_service_client()
            container_client = blob_service_client.get_container_client(BLOB_CONTAINER_NAME)
            
            # Buscar blobs que contengan el document_id en la carpeta propuestas/
            blobs = container_client.list_blobs(name_starts_with=PROPUESTAS_FOLDER)
            documento_encontrado = None
            
            for blob in blobs:
                if document_id in blob.name and blob.name.endswith('.docx'):
                    documento_encontrado = blob
                    break
            
            if documento_encontrado:
                # Generar URL pre-firmada con expiración de 2 horas
                url_presignada = generar_url_presignada(documento_encontrado.name, expiracion_minutos=120)  # 2 horas
                
                return func.HttpResponse(
                    json.dumps({
                        "message": "Documento encontrado",
                        "document_id": document_id,
                        "filename": documento_encontrado.name.replace(PROPUESTAS_FOLDER, ""),  # Remover prefijo para el filename
                        "url_presignada": url_presignada,
                        "expira_en_horas": 2,
                        "size_bytes": documento_encontrado.size,
                        "last_modified": documento_encontrado.last_modified.isoformat() if documento_encontrado.last_modified else None,
                        "status": "found"
                    }),
                    status_code=200,
                    mimetype="application/json"
                )
            else:
                return func.HttpResponse(
                    json.dumps({
                        "message": "Documento no encontrado",
                        "document_id": document_id,
                        "status": "not_found"
                    }),
                    status_code=404,
                    mimetype="application/json"
                )
                
        except Exception as search_error:
            return func.HttpResponse(
                json.dumps({
                    "error": f"Error buscando documento: {str(search_error)}",
                    "document_id": document_id,
                    "status": "error"
                }),
                status_code=500,
                mimetype="application/json"
            )
    
    except Exception as e:
        return func.HttpResponse(
            json.dumps({
                "error": f"Error interno del servidor: {str(e)}",
                "traceback": traceback.format_exc()
            }),
            status_code=500,
            mimetype="application/json"
        )

@app.function_name(name="listar_propuestas")
@app.route(route="listar_propuestas", methods=["GET"], auth_level=func.AuthLevel.ANONYMOUS)
def listar_propuestas(req: func.HttpRequest) -> func.HttpResponse:
    """Endpoint GET para listar todas las propuestas generadas"""
    try:
        blob_service_client = get_blob_service_client()
        container_client = blob_service_client.get_container_client(BLOB_CONTAINER_NAME)
        
        # Obtener todos los blobs de la carpeta propuestas/
        blobs = container_client.list_blobs(name_starts_with=PROPUESTAS_FOLDER)
        documentos = []
        
        for blob in blobs:
            if blob.name.endswith('.docx'):
                # Extraer información del nombre del archivo (sin el prefijo propuestas/)
                filename_without_prefix = blob.name.replace(PROPUESTAS_FOLDER, "")
                parts = filename_without_prefix.replace('.docx', '').split('_')
                
                # Generar URL pre-firmada para cada documento (1 hora de expiración)
                try:
                    url_presignada = generar_url_presignada(blob.name, expiracion_minutos=60)
                except Exception:
                    # Si falla la generación de SAS, usar URL pública (sin garantía de acceso)
                    url_presignada = f"https://{blob_service_client.account_name}.blob.core.windows.net/{BLOB_CONTAINER_NAME}/{blob.name}"
                
                documento_info = {
                    "filename": filename_without_prefix,
                    "full_blob_name": blob.name,
                    "url_presignada": url_presignada,
                    "expira_en_horas": 1,
                    "size_bytes": blob.size,
                    "last_modified": blob.last_modified.isoformat() if blob.last_modified else None,
                    "document_id": parts[2] if len(parts) >= 3 else "unknown",
                    "timestamp": parts[3] if len(parts) >= 4 else "unknown"
                }
                
                documentos.append(documento_info)
        
        # Ordenar por fecha de modificación (más reciente primero)
        documentos.sort(key=lambda x: x['last_modified'] or '', reverse=True)
        
        return func.HttpResponse(
            json.dumps({
                "message": "Listado de propuestas",
                "total_documentos": len(documentos),
                "documentos": documentos
            }),
            status_code=200,
            mimetype="application/json"
        )
        
    except Exception as e:
        return func.HttpResponse(
            json.dumps({
                "error": f"Error listando propuestas: {str(e)}",
                "traceback": traceback.format_exc()
            }),
            status_code=500,
            mimetype="application/json"
        )