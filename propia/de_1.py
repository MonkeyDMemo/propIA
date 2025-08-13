import azure.functions as func
import os
import re
import requests
from docx import Document
from docx.shared import Pt
from datetime import datetime, timedelta
from io import BytesIO
import json
import traceback
import logging
from azure.storage.blob import BlobServiceClient, generate_blob_sas, BlobSasPermissions

# Configuración de Azure OpenAI
AZURE_OPENAI_ENDPOINT = os.getenv("AZURE_OPENAI_ENDPOINT", "https://chabot-inventario-talento-aistudio.openai.azure.com/")
AZURE_OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")
DEPLOYMENT_NAME = os.getenv("DEPLOYMENT_NAME", "gpt-4o-mini")
API_VERSION = os.getenv("API_VERSION", "2024-02-15-preview")

# Configuración de Azure Storage
AZURE_STORAGE_CONNECTION_STRING = os.getenv("STORAGE_CONNECTION_STRING")
PLANTILLA_CONTAINER = "propia"
PROPUESTAS_CONTAINER = "propia"
PLANTILLA_BLOB_NAME = "plantilla/Plantilla-Propuesta.docx"
PROPUESTAS_FOLDER = "propuestas/"

# ========== FUNCIONES AUXILIARES ==========

def call_azure_openai(messages, max_tokens=1000):
    """Función para llamar a Azure OpenAI"""
    try:
        api_url = f"{AZURE_OPENAI_ENDPOINT}openai/deployments/{DEPLOYMENT_NAME}/chat/completions?api-version={API_VERSION}"

        headers = {
            "Content-Type": "application/json",
            "api-key": AZURE_OPENAI_API_KEY
        }

        data = {
            "messages": messages,
            "max_tokens": max_tokens
        }

        response = requests.post(api_url, headers=headers, json=data)

        if response.status_code == 200:
            return response.json()['choices'][0]['message']['content'].strip()
        else:
            logging.error(f"Error en Azure OpenAI: {response.status_code}, {response.text}")
            return None

    except Exception as e:
        logging.error(f"Error llamando a Azure OpenAI: {traceback.format_exc()}")
        return None

def limpiar_formato_markdown_mejorado(texto):
    """Limpia formato Markdown pero preserva información tabular y numérica"""
    # Convertir tablas markdown a formato de texto estructurado
    tabla_pattern = r'\|([^|]+\|)+\n\|[-:|\s]+\|\n(\|([^|]+\|)+\n?)+'
    tablas = re.findall(tabla_pattern, texto)
    
    for i, tabla in enumerate(tablas):
        lineas_tabla = tabla[0].split('\n')
        texto_tabla_formateado = ""
        
        for linea in lineas_tabla:
            if '|' in linea and not re.match(r'\|[-:\s|]+\|', linea):
                celdas = [celda.strip() for celda in linea.split('|') if celda.strip()]
                if celdas:
                    texto_tabla_formateado += " • " + " - ".join(celdas) + "\n"
        
        texto = texto.replace(tabla[0], texto_tabla_formateado)
    
    # Remover headers markdown
    texto = re.sub(r'^#{1,6}\s+', '', texto, flags=re.MULTILINE)
    # Remover bold markdown
    texto = re.sub(r'\*\*(.*?)\*\*', r'\1', texto)
    # Remover italic markdown
    texto = re.sub(r'\*(.*?)\*', r'\1', texto)
    # Convertir listas markdown en formato bullet
    texto = re.sub(r'^\s*[-\*\+]\s+', '• ', texto, flags=re.MULTILINE)
    # Remover bloques de código
    texto = re.sub(r'```(?:\w+)?\n?(.*?)\n?```', r'\1', texto, flags=re.DOTALL)
    # Limpiar líneas múltiples vacías
    texto = re.sub(r'\n\s*\n\s*\n', '\n\n', texto)
    
    return texto.strip()

# ========== FUNCIONES DE GENERACIÓN DE CONTENIDO ==========

def generar_resumen_ejecutivo(prompt_completo):
    """Genera contenido específico para resumen ejecutivo"""
    messages = [
        {
            "role": "system",
            "content": """Eres un consultor experto en propuestas técnicas. Genera ÚNICAMENTE un resumen ejecutivo profesional para documentos Word.

IMPORTANTE:
- NO uses formato Markdown (sin #, **, -, etc.)
- Escribe en párrafos corridos para Word
- Estilo profesional y ejecutivo
- Máximo 4 párrafos
- No incluyas títulos ni encabezados
- Enfócate solo en el resumen ejecutivo del proyecto"""
        },
        {
            "role": "user",
            "content": f"Basándote en esta información, genera únicamente el resumen ejecutivo profesional:\n\n{prompt_completo}"
        }
    ]
    
    contenido = call_azure_openai(messages, max_tokens=600)
    return limpiar_formato_markdown_mejorado(contenido) if contenido else None

def generar_alcance_minimo(prompt_completo):
    """Genera contenido específico para alcance mínimo"""
    messages = [
        {
            "role": "system",
            "content": """Eres un consultor experto en propuestas técnicas. Genera ÚNICAMENTE el alcance mínimo del proyecto para documentos Word.

IMPORTANTE:
- NO uses formato Markdown (sin #, **, -, etc.)
- Escribe en párrafos corridos para Word
- Describe qué incluye el proyecto específicamente
- Máximo 5 párrafos
- No incluyas títulos ni encabezados
- Enfócate solo en el alcance del proyecto"""
        },
        {
            "role": "user",
            "content": f"Basándote en esta información, genera únicamente el alcance mínimo del proyecto:\n\n{prompt_completo}"
        }
    ]
    
    contenido = call_azure_openai(messages, max_tokens=700)
    return limpiar_formato_markdown_mejorado(contenido) if contenido else None

def generar_plan_trabajo(prompt_completo):
    """Genera contenido específico para plan de trabajo"""
    messages = [
        {
            "role": "system",
            "content": """Eres un consultor experto en propuestas técnicas. Genera ÚNICAMENTE el plan de trabajo del proyecto para documentos Word.

IMPORTANTE:
- Si hay información tabular, preséntala usando bullets (•) en formato estructurado
- Conserva TODOS los números, fechas, porcentajes y datos cuantitativos
- Si hay fases con duraciones, inclúyelas con sus tiempos específicos
- Usa formato: "• Fase X: Descripción - Duración: X semanas - Porcentaje: X%"
- NO uses formato Markdown tabla (|---|) 
- Escribe en párrafos y listas con bullets para Word
- Máximo 6 párrafos o secciones con bullets
- Incluye TODA la información numérica disponible"""
        },
        {
            "role": "user",
            "content": f"Basándote en esta información, genera la descripción del plan de trabajo incluyendo TODAS las fases, duraciones y porcentajes mostrados:\n\n{prompt_completo}"
        }
    ]
    
    contenido = call_azure_openai(messages, max_tokens=800)
    return limpiar_formato_markdown_mejorado(contenido) if contenido else None

def generar_estructura_equipo(prompt_completo):
    """Genera contenido específico para estructura del equipo"""
    messages = [
        {
            "role": "system",
            "content": """Eres un consultor experto en propuestas técnicas. Genera ÚNICAMENTE la descripción del equipo de trabajo para documentos Word.

IMPORTANTE:
- Si hay información de roles con costos, preséntala usando bullets (•)
- Conserva TODOS los números, tarifas, horas y subtotales
- Usa formato: "• Rol: Descripción - Dedicación: X% - Horas: X - Tarifa: $X - Subtotal: $X"
- NO uses formato Markdown tabla (|---|)
- Incluye descuentos y totales si están disponibles
- Escribe en párrafos y listas con bullets para Word
- Máximo 5 párrafos o secciones con bullets
- Incluye TODA la información numérica y financiera disponible"""
        },
        {
            "role": "user",
            "content": f"Basándote en esta información, genera la descripción del equipo incluyendo TODOS los roles, costos, tarifas y totales mostrados:\n\n{prompt_completo}"
        }
    ]
    
    contenido = call_azure_openai(messages, max_tokens=800)
    return limpiar_formato_markdown_mejorado(contenido) if contenido else None

def generar_inversion_detallada(prompt_completo):
    """Genera contenido específico para inversión detallada"""
    messages = [
        {
            "role": "system",
            "content": """Eres un consultor experto en propuestas técnicas. Genera ÚNICAMENTE la explicación de la inversión detallada para documentos Word.

IMPORTANTE:
- Conserva TODOS los números, montos, porcentajes y cifras exactas
- Si hay desglose de costos, preséntalo usando bullets (•)
- Usa formato: "• Concepto: Descripción - Monto: $X,XXX MXN"
- NO uses formato Markdown tabla (|---|)
- Incluye servicios profesionales, costos de setup, costos operativos
- Menciona la inversión total inicial con el monto exacto
- Escribe en párrafos y listas con bullets para Word
- Máximo 4 párrafos o secciones con bullets
- Incluye TODA la información financiera disponible"""
        },
        {
            "role": "user",
            "content": f"Basándote en esta información, genera la explicación de la inversión incluyendo TODOS los montos, costos y totales mostrados:\n\n{prompt_completo}"
        }
    ]
    
    contenido = call_azure_openai(messages, max_tokens=700)
    return limpiar_formato_markdown_mejorado(contenido) if contenido else None

def generar_supuestos_condiciones(prompt_completo):
    """Genera contenido específico para supuestos y condiciones"""
    messages = [
        {
            "role": "system",
            "content": """Eres un consultor experto en propuestas técnicas. Genera ÚNICAMENTE los supuestos y condiciones del proyecto para documentos Word.

IMPORTANTE:
- NO uses formato Markdown (sin #, **, -, etc.)
- Escribe en párrafos corridos para Word
- Describe supuestos técnicos y condiciones comerciales
- Máximo 4 párrafos
- No incluyas títulos ni encabezados
- Enfócate en aspectos clave del proyecto"""
        },
        {
            "role": "user",
            "content": f"Basándote en esta información, genera únicamente los supuestos y condiciones:\n\n{prompt_completo}"
        }
    ]
    
    contenido = call_azure_openai(messages, max_tokens=700)
    return limpiar_formato_markdown_mejorado(contenido) if contenido else None

def generar_contenido_generico(prompt_completo, placeholder):
    """Función genérica para contenido personalizado"""
    messages = [
        {
            "role": "system",
            "content": """Eres un consultor experto en propuestas técnicas. Genera contenido profesional para documentos Word.

IMPORTANTE:
- Conserva TODOS los números, fechas, porcentajes, montos y datos cuantitativos
- Si hay información tabular, conviértela a bullets (•) estructurados
- Usa formato: "• Elemento: Descripción - Datos: valores específicos"
- NO uses formato Markdown tabla (|---|) ni headers (#)
- Escribe en párrafos corridos y listas con bullets para Word
- Usa un estilo profesional y ejecutivo
- Interpreta la información y redacta de forma natural
- Incluye TODA la información numérica disponible
- Genera contenido fluido sin perder datos específicos"""
        },
        {
            "role": "user",
            "content": f"Genera contenido para la sección {placeholder} basándote en esta información, conservando TODOS los números y datos específicos:\n\n{prompt_completo}"
        }
    ]
    
    contenido = call_azure_openai(messages, max_tokens=1000)
    return limpiar_formato_markdown_mejorado(contenido) if contenido else None

# ========== FUNCIONES DE PROCESAMIENTO DE DOCUMENTOS ==========

def set_font_format(run, font_name="Arial Nova Cond", font_size=11.5):
    """Establece el formato de fuente para un run específico"""
    run.font.name = font_name
    run.font.size = Pt(font_size)

def replace_in_paragraph(paragraph, old_text, new_text, font_name="Arial Nova Cond", font_size=11.5):
    """Reemplaza texto en párrafos"""
    if old_text in paragraph.text:
        full_text = paragraph.text.replace(old_text, new_text)
        
        # Limpiar runs existentes
        for run in paragraph.runs:
            run.clear()
        
        # Crear nuevo run con formato
        new_run = paragraph.add_run(full_text)
        set_font_format(new_run, font_name, font_size)
        
        return True
    return False

def replace_in_tables(doc, old_text, new_text, font_name="Arial Nova Cond", font_size=12):
    """Reemplaza texto en tablas del documento"""
    replacements_made = 0
    
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if replace_in_paragraph(paragraph, old_text, new_text, font_name, font_size):
                        replacements_made += 1
    
    return replacements_made

def replace_in_textboxes(doc, old_text, new_text):
    """Reemplaza texto en cuadros de texto usando el método que funciona"""
    replacements_made = 0
    
    try:
        # Buscar elementos p que contienen exactamente el placeholder
        for element in doc.element.iter():
            if element.tag.endswith('}p'):
                # Obtener el texto completo del párrafo concatenando todos los elementos t
                full_text = ""
                text_elements = []
                
                for child in element.iter():
                    if child.tag.endswith('}t') and child.text:
                        full_text += child.text
                        text_elements.append(child)
                
                # Si el texto completo contiene el placeholder
                if old_text in full_text:
                    # Verificar si está en txbxContent (textbox content)
                    parent = element.getparent()
                    if parent is not None and 'txbxContent' in parent.tag:
                        # Limpiar todos los elementos de texto existentes
                        for text_elem in text_elements:
                            text_elem.text = ""
                        
                        # Poner el texto reemplazado en el primer elemento
                        if text_elements:
                            new_text_content = full_text.replace(old_text, new_text)
                            text_elements[0].text = new_text_content
                            replacements_made += 1
        
        # Método de respaldo
        if replacements_made == 0:
            for element in doc.element.body.iter():
                for run_element in element.iter():
                    if run_element.tag.endswith('}r'):
                        for text_element in run_element.iter():
                            if text_element.tag.endswith('}t') and text_element.text:
                                if old_text in text_element.text:
                                    text_element.text = text_element.text.replace(old_text, new_text)
                                    replacements_made += 1
        
    except Exception as e:
        logging.error(f"Error procesando cuadros de texto: {e}")
    
    return replacements_made

# ========== FUNCIONES DE EXTRACCIÓN Y UTILIDADES ==========

def extraer_informacion_empresa(prompt_completo):
    """Extrae información específica de la empresa del prompt"""
    empresa_pattern = r'(?:para\s+|de\s+|empresa\s+)([A-Z][a-zA-Z\s&]+(?:SA\s+de\s+CV|S\.A\.|Inc\.|Corp\.|Ltd\.)?)'
    fecha_pattern = r'(\d{1,2}\s+de\s+\w+\s+de\s+\d{4}|\d{1,2}/\d{1,2}/\d{4})'
    titulo_pattern = r'#\s*(.*?)(?:\n|$)'
    
    info = {
        'empresa': 'Cliente Estimado',
        'fecha': datetime.now().strftime('%d de %B de %Y'),
        'titulo': 'Propuesta Técnica'
    }
    
    # Extraer nombre de empresa
    empresa_match = re.search(empresa_pattern, prompt_completo, re.IGNORECASE)
    if empresa_match:
        info['empresa'] = empresa_match.group(1).strip()
    
    # Extraer fecha
    fecha_match = re.search(fecha_pattern, prompt_completo)
    if fecha_match:
        info['fecha'] = fecha_match.group(1)
    
    # Extraer título del documento
    titulo_match = re.search(titulo_pattern, prompt_completo)
    if titulo_match:
        info['titulo'] = titulo_match.group(1).strip()
    
    return info

def generar_carta_presentacion(prompt_completo):
    """Genera contenido específico para carta de presentación"""
    info_empresa = extraer_informacion_empresa(prompt_completo)
    
    messages = [
        {
            "role": "system",
            "content": f"""Eres un consultor experto en propuestas técnicas. Genera ÚNICAMENTE una carta de presentación profesional para documentos Word.

IMPORTANTE:
- NO uses formato Markdown (sin #, **, -, etc.)
- Escribe en párrafos corridos para Word
- Usa un tono profesional y cordial
- Menciona específicamente el proyecto y sus características principales
- Incluye los datos extraídos: Empresa: {info_empresa['empresa']}, Fecha: {info_empresa['fecha']}
- Máximo 3 párrafos
- No incluyas títulos ni encabezados
- Enfócate en el valor que HITSS puede aportar al proyecto"""
        },
        {
            "role": "user",
            "content": f"Genera una carta de presentación profesional para la empresa {info_empresa['empresa']} basándote en este proyecto:\n\n{prompt_completo}"
        }
    ]
    
    contenido = call_azure_openai(messages, max_tokens=500)
    
    if contenido:
        carta_personalizada = f"""Estimado Equipo,

Ciudad de México, México {info_empresa['fecha']}

En representación de HITSS, agradecemos profundamente la oportunidad que nos brindan de presentar nuestra propuesta para {info_empresa['titulo'].lower()}. {contenido}

Atentamente,
Sergio Portales Aburto"""
        
        return limpiar_formato_markdown_mejorado(carta_personalizada)
    
    return None

def generar_titulo_fecha(prompt_completo):
    """Genera título y fecha para el cuadro de texto de la página 1"""
    info_empresa = extraer_informacion_empresa(prompt_completo)
    
    titulo_fecha_content = f"""{info_empresa['titulo']}
{info_empresa['fecha']}"""
    
    return titulo_fecha_content

# ========== FUNCIONES DE AZURE STORAGE ==========

def descargar_plantilla():
    """Descarga la plantilla desde Azure Blob Storage"""
    try:
        blob_service_client = BlobServiceClient.from_connection_string(AZURE_STORAGE_CONNECTION_STRING)
        blob_client = blob_service_client.get_blob_client(container=PLANTILLA_CONTAINER, blob=PLANTILLA_BLOB_NAME)
        
        # Descargar el blob en memoria
        plantilla_stream = BytesIO()
        blob_data = blob_client.download_blob()
        plantilla_stream.write(blob_data.readall())
        plantilla_stream.seek(0)
        
        return plantilla_stream
    except Exception as e:
        logging.error(f"Error descargando plantilla: {traceback.format_exc()}")
        return None

def subir_a_blob_storage(nombre_archivo, contenido):
    """Sube archivo al Blob Storage"""
    try:
        blob_service_client = BlobServiceClient.from_connection_string(AZURE_STORAGE_CONNECTION_STRING)
        # Añadir la carpeta propuestas/ al nombre del blob
        blob_name = f"{PROPUESTAS_FOLDER}{nombre_archivo}"
        blob_client = blob_service_client.get_blob_client(container=PROPUESTAS_CONTAINER, blob=blob_name)
        
        contenido.seek(0)
        blob_client.upload_blob(contenido, overwrite=True)
        
        return f"https://{blob_client.account_name}.blob.core.windows.net/{PROPUESTAS_CONTAINER}/{blob_name}"
    except Exception as e:
        logging.error(f"Error subiendo archivo: {traceback.format_exc()}")
        return None

def generar_url_presignada(nombre_archivo, expiracion_minutos=60):
    """Genera una URL pre-firmada (SAS) para acceder al archivo"""
    try:
        blob_service_client = BlobServiceClient.from_connection_string(AZURE_STORAGE_CONNECTION_STRING)
        # Añadir la carpeta propuestas/ al nombre del blob
        blob_name = f"{PROPUESTAS_FOLDER}{nombre_archivo}"
        blob_client = blob_service_client.get_blob_client(container=PROPUESTAS_CONTAINER, blob=blob_name)

        sas_token = generate_blob_sas(
            account_name=blob_client.account_name,
            container_name=PROPUESTAS_CONTAINER,
            blob_name=blob_name,
            account_key=blob_service_client.credential.account_key,
            permission=BlobSasPermissions(read=True),
            expiry=datetime.utcnow() + timedelta(minutes=expiracion_minutos)
        )

        url_presignada = f"https://{blob_client.account_name}.blob.core.windows.net/{PROPUESTAS_CONTAINER}/{blob_name}?{sas_token}"
        return url_presignada
    except Exception as e:
        logging.error(f"Error generando URL pre-firmada: {traceback.format_exc()}")
        return None

# ========== FUNCIÓN PRINCIPAL DE PROCESAMIENTO ==========

def procesar_propuesta_completa(prompt_completo, placeholders_personalizados=None):
    """Procesa una propuesta completa"""
    
    # Descargar plantilla
    plantilla_stream = descargar_plantilla()
    if not plantilla_stream:
        raise Exception("No se pudo descargar la plantilla")
    
    logging.info(f"Procesando propuesta: {len(prompt_completo)} caracteres")
    
    # Extraer información de la empresa
    info_empresa = extraer_informacion_empresa(prompt_completo)
    logging.info(f"Empresa: {info_empresa['empresa']}, Fecha: {info_empresa['fecha']}, Título: {info_empresa['titulo']}")
    
    # Definir placeholders con funciones de generación
    placeholders_config = {
        "[RESUMEN]": generar_resumen_ejecutivo,
        "[ALCANCE]": generar_alcance_minimo,
        "[PLAN_TRABAJO]": generar_plan_trabajo,
        "[EQUIPO]": generar_estructura_equipo,
        "[INVERSION]": generar_inversion_detallada,
        "[SUPUESTOS]": generar_supuestos_condiciones,
        "[CARTA_PRESENTACION]": generar_carta_presentacion,
        "[titulo]": lambda prompt: generar_titulo_fecha(prompt).split('\n')[0],
        "[fecha]": lambda prompt: generar_titulo_fecha(prompt).split('\n')[1]
    }
    
    # Agregar placeholders personalizados si se proporcionan
    if placeholders_personalizados:
        placeholders_config.update(placeholders_personalizados)
    
    # Cargar documento
    doc = Document(plantilla_stream)
    cambios_totales = 0
    
    # Procesar cada placeholder
    for placeholder, funcion_generadora in placeholders_config.items():
        logging.info(f"Procesando: {placeholder}")
        
        try:
            # Generar contenido específico
            contenido_generado = funcion_generadora(prompt_completo)
            
            if not contenido_generado:
                logging.warning(f"No se pudo generar contenido para {placeholder}")
                continue
            
            logging.info(f"Contenido generado: {len(contenido_generado)} caracteres")
            
            replacements_for_this_item = 0
            
            # 1. Buscar en párrafos normales
            for paragraph in doc.paragraphs:
                if replace_in_paragraph(paragraph, placeholder, contenido_generado):
                    replacements_for_this_item += 1
            
            # 2. Buscar en tablas
            table_replacements = replace_in_tables(doc, placeholder, contenido_generado)
            replacements_for_this_item += table_replacements
            
            # 3. Buscar en cuadros de texto
            textbox_replacements = replace_in_textboxes(doc, placeholder, contenido_generado)
            replacements_for_this_item += textbox_replacements
            
            if replacements_for_this_item > 0:
                logging.info(f"Realizados {replacements_for_this_item} reemplazos para {placeholder}")
                cambios_totales += replacements_for_this_item
            else:
                logging.warning(f"No se encontró el placeholder {placeholder} en el documento")
                
        except Exception as e:
            logging.error(f"Error procesando {placeholder}: {traceback.format_exc()}")
    
    if cambios_totales == 0:
        raise Exception("No se realizaron cambios en el documento")
    
    # Guardar documento en memoria
    output_stream = BytesIO()
    doc.save(output_stream)
    
    # Generar nombre de archivo
    timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
    empresa_clean = re.sub(r'[^\w\s-]', '', info_empresa['empresa']).strip()[:20]
    nombre_archivo = f"Propuesta_{empresa_clean}_{timestamp}.docx"
    
    # Subir a Blob Storage
    url_archivo = subir_a_blob_storage(nombre_archivo, output_stream)
    if not url_archivo:
        raise Exception("Error al subir el archivo a Azure Blob Storage")
    
    logging.info(f"Documento guardado: {nombre_archivo}")
    logging.info(f"Total de reemplazos realizados: {cambios_totales}")
    
    # Generar URL pre-firmada
    url_presignada = generar_url_presignada(nombre_archivo)
    
    return url_presignada

# ========== FUNCIÓN PRINCIPAL DE AZURE FUNCTION ==========

def main(req: func.HttpRequest) -> func.HttpResponse:
    try:
        logging.info("Inicio de la función de generación de propuestas")
        
        # Intentar obtener el contenido de diferentes formas
        prompt_completo = None
        placeholders_personalizados = None
        
        # Primero intentar obtener JSON
        try:
            req_body = req.get_json()
            if req_body:
                prompt_completo = req_body.get('prompt')
                placeholders_personalizados = req_body.get('placeholders_personalizados', None)
        except ValueError:
            # Si no es JSON, intentar obtener como texto plano
            logging.info("No se recibió JSON válido, intentando como texto plano")
            req_body_text = req.get_body().decode('utf-8').strip()
            if req_body_text:
                prompt_completo = req_body_text
        
        # Si no se obtuvo prompt de ninguna forma, verificar parámetros de query
        if not prompt_completo:
            prompt_completo = req.params.get('prompt')
        
        # Validar que se recibió el prompt
        if not prompt_completo:
            return func.HttpResponse(
                json.dumps({
                    "error": "Se requiere un prompt",
                    "instrucciones": "Envía el prompt como: 1) Campo 'prompt' en JSON body, 2) Texto plano en el body, o 3) Parámetro 'prompt' en la URL"
                }),
                status_code=400,
                mimetype='application/json'
            )
        
        # Si se proporcionan placeholders personalizados, convertirlos en funciones
        if placeholders_personalizados:
            placeholders_funciones = {}
            for placeholder, contenido in placeholders_personalizados.items():
                # Si el contenido es string, crear función que lo retorne
                if isinstance(contenido, str):
                    placeholders_funciones[placeholder] = lambda prompt, c=contenido: c
                else:
                    # Si es otro tipo, usar función genérica
                    placeholders_funciones[placeholder] = lambda prompt, p=placeholder: generar_contenido_generico(prompt, p)
            placeholders_personalizados = placeholders_funciones
        
        logging.info(f"Procesando propuesta con prompt de {len(prompt_completo)} caracteres")
        
        # Procesar la propuesta
        url_presignada = procesar_propuesta_completa(prompt_completo, placeholders_personalizados)
        
        if not url_presignada:
            return func.HttpResponse(
                json.dumps({"error": "Error generando la propuesta"}),
                status_code=500,
                mimetype='application/json'
            )
        
        # Extraer información de la empresa para la respuesta
        info_empresa = extraer_informacion_empresa(prompt_completo)
        
        response_data = {
            "url": url_presignada,
            "empresa": info_empresa['empresa'],
            "fecha": info_empresa['fecha'],
            "titulo": info_empresa['titulo'],
            "mensaje": "Propuesta generada exitosamente"
        }
        
        logging.info(f"Propuesta generada exitosamente: {url_presignada}")
        
        return func.HttpResponse(
            json.dumps(response_data),
            status_code=200,
            mimetype='application/json'
        )
        
    except Exception as e:
        logging.error(f"Error en la función: {traceback.format_exc()}")
        
        return func.HttpResponse(
            json.dumps({
                "error": "Error interno del servidor",
                "detalle": str(e),
                "traceback": traceback.format_exc()
            }),
            status_code=500,
            mimetype='application/json'
        )